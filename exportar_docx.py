import sqlite3
import pandas as pd
import sys
import datetime
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from exportar_pdf import (
    CRONOGRAMA_ID_FORZADO,
    select_cronograma_id,
    obtener_datos_cronograma,
    procesar_guardias,
    construir_semanas,
    simplificar_nombre,
    MESES
)

def set_cell_background(cell, fill_hex):
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def style_header_cell(cell, text):
    set_cell_background(cell, "1C3144")
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.font.name = 'Helvetica'
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def generar_docx(crono_id, fecha_inicio, fecha_fin, notas, guardias_por_fecha, semanas, output_filename):
    doc = Document()
    
    # Configurar página en A4 Horizontal (Landscape)
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    
    # Márgenes de 0.5 pulgadas
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    
    # Título del documento
    fecha_dt = datetime.datetime.strptime(fecha_inicio, "%Y-%m-%d")
    mes_str = MESES[fecha_dt.month]
    anio = fecha_dt.year
    
    title_text = "CRONOGRAMA DE GUARDIAS - ÁREA MÉDICA UTI"
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    run_title = p_title.add_run(title_text)
    run_title.font.name = 'Helvetica'
    run_title.font.size = Pt(15)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0x1C, 0x31, 0x44)
    
    subtitle_text = f"PERÍODO: {mes_str} {anio}"
    if notas and notas != "Generado via CLI":
        subtitle_text += f" ({notas})"
        
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(8)
    run_sub = p_sub.add_run(subtitle_text)
    run_sub.font.name = 'Helvetica'
    run_sub.font.size = Pt(10)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    # Crear Tabla (1 fila de encabezado + filas de semanas)
    headers = ['LUNES', 'MARTES', 'MIÉRCOLES', 'JUEVES', 'VIERNES', 'SÁBADO', 'DOMINGO']
    num_rows = len(semanas) + 1
    table = doc.add_table(rows=num_rows, cols=7)
    table.style = 'Table Grid'
    
    # Ancho de columnas: distribuir el espacio horizontal de la página A4 (aprox 10.69 pulgadas útiles)
    col_width = Inches(1.52) # 1.52 * 7 = 10.64 pulgadas
    
    # 1. Estilar el encabezado
    hdr_cells = table.rows[0].cells
    for i, h_text in enumerate(headers):
        style_header_cell(hdr_cells[i], h_text)
        hdr_cells[i].width = col_width
        
    # Set height of header row
    table.rows[0].height = Pt(20)
    
    # 2. Llenar las semanas
    for r_idx, week in enumerate(semanas):
        row = table.rows[r_idx + 1]
        # Altura estimada mínima de fila en Word
        row.height = Pt(75)
        
        for c_idx, day_info in enumerate(week):
            cell = row.cells[c_idx]
            cell.width = col_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            
            # Sombreado de fines de semana/feriados
            if day_info["is_shaded"]:
                set_cell_background(cell, "E8E8E8")
                
            if not day_info["in_range"]:
                # Celda fuera de rango (mes anterior o posterior en la grilla)
                continue
                
            # Escribir el número del día (alineado a la derecha)
            p_day = cell.paragraphs[0]
            p_day.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p_day.paragraph_format.space_after = Pt(2)
            run_day = p_day.add_run(day_info["day_num"])
            run_day.font.name = 'Helvetica'
            run_day.font.size = Pt(9)
            run_day.font.bold = True
            run_day.font.color.rgb = RGBColor(0x1C, 0x31, 0x44)
            
            # Obtener datos de asignaciones para este día
            data_dia = guardias_por_fecha.get(day_info["date_str"], {})
            if data_dia:
                # Escribir Guardias de 24 hs
                if data_dia.get("Guardia"):
                    for n in data_dia["Guardia"]:
                        p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.line_spacing = 1.05
                        run = p.add_run(f"- {simplificar_nombre(n)}")
                        run.font.name = 'Helvetica'
                        run.font.size = Pt(8.5)
                        
                # Escribir Turnos de Día (8-20)
                if data_dia.get("Día"):
                    p_title = cell.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_title.paragraph_format.space_after = Pt(1)
                    p_title.paragraph_format.line_spacing = 1.05
                    run = p_title.add_run("Día (8-20):")
                    run.font.name = 'Helvetica'
                    run.font.size = Pt(7.5)
                    run.font.bold = True
                    
                    for n in data_dia["Día"]:
                        p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.line_spacing = 1.05
                        run = p.add_run(f"- {simplificar_nombre(n)}")
                        run.font.name = 'Helvetica'
                        run.font.size = Pt(8.5)
                        
                # Escribir Turnos de Noche (20-8)
                if data_dia.get("Noche"):
                    p_title = cell.add_paragraph()
                    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p_title.paragraph_format.space_after = Pt(1)
                    p_title.paragraph_format.line_spacing = 1.05
                    run = p_title.add_run("Noche (20-8):")
                    run.font.name = 'Helvetica'
                    run.font.size = Pt(7.5)
                    run.font.bold = True
                    
                    for n in data_dia["Noche"]:
                        p = cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        p.paragraph_format.space_after = Pt(1)
                        p.paragraph_format.line_spacing = 1.05
                        run = p.add_run(f"- {simplificar_nombre(n)}")
                        run.font.name = 'Helvetica'
                        run.font.size = Pt(8.5)
                        
    doc.save(output_filename)

def main():
    if CRONOGRAMA_ID_FORZADO is not None:
        crono_id = CRONOGRAMA_ID_FORZADO
    elif len(sys.argv) > 1:
        try:
            crono_id = int(sys.argv[1])
        except ValueError:
            print("ERROR: El ID del cronograma debe ser un número entero.")
            return
    else:
        crono_id = select_cronograma_id()
        if crono_id is None:
            return
            
    print(f"\nProcesando Cronograma ID: {crono_id} para exportar a Word (.docx)...")
    datos = obtener_datos_cronograma(crono_id)
    if not datos:
        return
        
    fecha_inicio = datos["fecha_inicio"]
    fecha_fin = datos["fecha_fin"]
    notas = datos["notas"]
    guardias_rows = datos["guardias"]
    
    if not guardias_rows:
        print(f"\nADVERTENCIA: No se encontraron asignaciones del Área Médica (Servicio 3) para el Cronograma ID {crono_id}.")
        print("¿Desea exportar el Word de todos modos? Se generará el calendario vacío.")
        try:
            cont = input("¿Continuar? (s/n): ").strip().lower()
            if cont != 's':
                print("Operación cancelada.")
                return
        except (KeyboardInterrupt, ValueError):
            return
            
    print(f"[OK] Cargadas {len(guardias_rows)} asignaciones de guardia.")
    print(f"     Período: {fecha_inicio} al {fecha_fin}")
    
    # Procesar datos
    guardias_por_fecha = procesar_guardias(guardias_rows)
    semanas = construir_semanas(fecha_inicio, fecha_fin)
    
    # Nombre del archivo final
    fecha_dt = datetime.datetime.strptime(fecha_inicio, "%Y-%m-%d")
    mes_str = MESES[fecha_dt.month].capitalize()
    anio = fecha_dt.year
    output_filename = f"Cronograma_Medicos_{mes_str}_{anio}.docx"
    
    print(f"Generando Word (.docx): {output_filename}...")
    try:
        generar_docx(crono_id, fecha_inicio, fecha_fin, notas, guardias_por_fecha, semanas, output_filename)
        print(f"\n[SUCCESS] Word (.docx) generado con éxito: {output_filename}!")
    except Exception as e:
        print(f"\n[ERROR] Ocurrió un error al generar el Word: {e}")

if __name__ == '__main__':
    main()
